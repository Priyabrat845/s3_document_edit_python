import boto3
#TEST
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
def insert_paragraphs_and_tables_after_paragraph_number(paragraph_number, paragraph_list, table_data_list):
    zip_data = zip(paragraph_list,table_data_list)
    output_file = '/tmp/output.docx'

    # Read Word document from S3
    client = boto3.client('s3')
    bucket_name = ''
    file_key = ''
    output_file_key = ''
    response = client.get_object(Bucket=bucket_name, Key=file_key)
    document_data = response['Body'].read()

    # Load the input document
    doc = Document(document_data)


    # Iterate over the paragraphs in the document
    for i, paragraph in enumerate(doc.paragraphs):
        # Check if the current paragraph is the desired paragraph number
        if i + 1 == paragraph_number:
            # Insert new paragraphs after the current paragraph
            for paragraph_text,table_data in zip_data:
                new_paragraph = doc.add_paragraph()
                new_paragraph.text = paragraph_text
                print('Paragraph inserted')
                # Insert new tables after the current paragraph
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))

                for row_index, row_data in enumerate(table_data):
                    for col_index, cell_data in enumerate(row_data):
                        cell = table.cell(row_index, col_index)
                        cell.text = cell_data

                        # Bold the table headers
                        if row_index == 0:
                          paragraph = cell.paragraphs[0]
                          #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                          run = paragraph.runs[0]
                          run.font.size = Pt(12)
                          run.bold = True



                table.style = 'Table Grid'
                table.autofit = False
                table.width = Cm(10)
                print('Table inserted')

    # Save the output document
    doc.save(output_file)

    # Upload the modified document back to S3
    with open(output_file, 'rb') as file:
      s3.upload_fileobj(file, bucket_name, output_file_key)





# Example usage
paragraph_number = 10
paragraph_list = ['Table - 1','Table - 2','Table - 3']
insert_paragraphs_and_tables_after_paragraph_number(paragraph_number, paragraph_list, table_data_list)